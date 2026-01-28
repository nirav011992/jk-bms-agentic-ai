"""Utility functions for processing different document types."""
import io
from typing import Tuple, List
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


class TextChunker:
    """Intelligently chunk text while preserving sentence boundaries."""

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 3500,
        overlap: int = 200,
    ) -> List[str]:
        """
        Split text into chunks while preserving sentence boundaries.

        Args:
            text: Text to chunk
            chunk_size: Target characters per chunk (3000-4000 recommended)
            overlap: Overlapping characters between chunks for context

        Returns:
            List of text chunks
        """
        if not text or len(text.strip()) == 0:
            return []

        # Clean text
        text = text.strip()

        # If text is smaller than chunk size, return as single chunk
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        current_pos = 0

        while current_pos < len(text):
            # Define chunk boundaries
            chunk_end = min(current_pos + chunk_size, len(text))

            # Try to find a sentence boundary near the end
            if chunk_end < len(text):
                # Look backwards for sentence-ending punctuation
                search_text = text[current_pos:chunk_end]
                last_sentence_end = max(
                    search_text.rfind('. '),
                    search_text.rfind('! '),
                    search_text.rfind('? '),
                )

                if last_sentence_end != -1:
                    # Found a sentence boundary, use it
                    chunk_end = current_pos + last_sentence_end + 2
                else:
                    # No sentence boundary found, look for paragraph break
                    last_para_end = search_text.rfind('\n\n')
                    if last_para_end != -1:
                        chunk_end = current_pos + last_para_end + 2
                    # else: use the hard chunk_size limit

            # Extract chunk
            chunk = text[current_pos:chunk_end].strip()
            if chunk:
                chunks.append(chunk)

            # Move position forward (with overlap for context)
            if chunk_end >= len(text):
                break

            current_pos = chunk_end - overlap

        return chunks


class DocumentProcessor:
    """Process different document formats and extract text content."""

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing null bytes and normalizing whitespace.

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text safe for PostgreSQL UTF-8 storage
        """
        # Remove null bytes that cause PostgreSQL UTF-8 errors
        text = text.replace('\x00', '')

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive blank lines (more than 2 consecutive)
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')

        return text.strip()

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file.

        Args:
            file_content: Raw bytes of the PDF file

        Returns:
            Extracted text content as string

        Raises:
            ValueError: If PDF cannot be processed
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)

            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

            extracted_text = "\n".join(text_content)

            if not extracted_text.strip():
                raise ValueError("No text content found in PDF")

            # Clean text to remove null bytes and normalize
            return DocumentProcessor.clean_text(extracted_text)

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file.

        Args:
            file_content: Raw bytes of the DOCX file

        Returns:
            Extracted text content as string

        Raises:
            ValueError: If DOCX cannot be processed
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            extracted_text = "\n".join(text_content)

            if not extracted_text.strip():
                raise ValueError("No text content found in DOCX")

            # Clean text to remove null bytes and normalize
            return DocumentProcessor.clean_text(extracted_text)

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from plain text file.

        Args:
            file_content: Raw bytes of the text file

        Returns:
            Decoded text content as string

        Raises:
            ValueError: If text cannot be decoded
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1 which accepts all byte values
                text = file_content.decode('latin-1')

            # Clean text to remove null bytes and normalize
            return DocumentProcessor.clean_text(text)

        except Exception as e:
            raise ValueError(f"Failed to decode text file: {str(e)}")

    @staticmethod
    def process_document(file_content: bytes, filename: str, content_type: str = None) -> Tuple[str, str]:
        """Process document based on file type and extract text.

        Args:
            file_content: Raw bytes of the file
            filename: Name of the file
            content_type: MIME type of the file (optional)

        Returns:
            Tuple of (extracted_text, detected_format)

        Raises:
            ValueError: If file format is not supported or processing fails
        """
        filename_lower = filename.lower()

        # Determine file type
        if filename_lower.endswith('.pdf') or content_type == 'application/pdf':
            text = DocumentProcessor.extract_text_from_pdf(file_content)
            return text, 'pdf'

        elif filename_lower.endswith('.docx') or content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]:
            text = DocumentProcessor.extract_text_from_docx(file_content)
            return text, 'docx'

        elif filename_lower.endswith('.doc') or content_type == 'application/msword':
            # .doc files (old Word format) are not supported by python-docx
            raise ValueError("Legacy .doc files are not supported. Please convert to .docx or PDF")

        elif filename_lower.endswith(('.txt', '.md', '.markdown')) or \
             content_type in ['text/plain', 'text/markdown']:
            text = DocumentProcessor.extract_text_from_txt(file_content)
            return text, 'text'

        else:
            raise ValueError(
                f"Unsupported file format: {filename}. "
                "Supported formats: PDF, DOCX, TXT, MD"
            )


document_processor = DocumentProcessor()
